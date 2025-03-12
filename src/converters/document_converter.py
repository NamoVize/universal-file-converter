"""
Document Converter Module for the Universal File Converter
"""
import os
import logging
import subprocess
import csv
import codecs
import pandas as pd
from PIL import Image
import pytesseract
from docx import Document
from pdf2docx import Converter
import PyPDF2
from pptx import Presentation

from utils.file_utils import get_output_path


class DocumentConverter:
    """Handles conversion between various document formats"""
    
    # Supported formats and their file extensions
    SUPPORTED_INPUT_FORMATS = {
        "pdf", "docx", "doc", "txt", "rtf", "odt", "xlsx", "xls", "csv", "pptx", "ppt"
    }
    
    SUPPORTED_OUTPUT_FORMATS = {
        "pdf", "docx", "txt", "rtf", "odt", "xlsx", "csv", "pptx", "html"
    }
    
    def __init__(self):
        """Initialize the document converter"""
        self.logger = logging.getLogger(__name__)
    
    def convert(self, input_path, output_format, output_dir, **options):
        """
        Convert a document file to the specified format
        
        Args:
            input_path (str): Path to the input document file
            output_format (str): Desired output format (e.g., "pdf", "docx")
            output_dir (str): Directory to save the output file
            **options: Additional options for conversion
                - overwrite (bool): Whether to overwrite existing files (default: False)
        
        Returns:
            bool: True if conversion was successful, False otherwise
        """
        try:
            # Validate input format
            input_extension = os.path.splitext(input_path)[1][1:].lower()
            if input_extension not in self.SUPPORTED_INPUT_FORMATS:
                self.logger.error(f"Unsupported input format: {input_extension}")
                return False
            
            # Validate output format
            output_format = output_format.lower().replace(".", "")
            if output_format not in self.SUPPORTED_OUTPUT_FORMATS:
                self.logger.error(f"Unsupported output format: {output_format}")
                return False
            
            # Get output path
            output_path = get_output_path(input_path, output_format, output_dir)
            
            # Check if output file exists and overwrite is not enabled
            if os.path.exists(output_path) and not options.get("overwrite", False):
                self.logger.warning(f"Output file already exists: {output_path}")
                return False
            
            # Handle conversion based on input and output format
            if input_extension == "pdf" and output_format == "docx":
                return self._convert_pdf_to_docx(input_path, output_path)
            
            elif input_extension == "pdf" and output_format == "txt":
                return self._convert_pdf_to_txt(input_path, output_path)
            
            elif input_extension in ["docx", "doc"] and output_format == "pdf":
                return self._convert_doc_to_pdf(input_path, output_path)
            
            elif input_extension in ["docx", "doc"] and output_format == "txt":
                return self._convert_doc_to_txt(input_path, output_path)
            
            elif input_extension == "txt" and output_format == "pdf":
                return self._convert_txt_to_pdf(input_path, output_path)
            
            elif input_extension == "txt" and output_format == "docx":
                return self._convert_txt_to_docx(input_path, output_path)
            
            elif input_extension in ["xlsx", "xls"] and output_format == "csv":
                return self._convert_excel_to_csv(input_path, output_path)
            
            elif input_extension == "csv" and output_format == "xlsx":
                return self._convert_csv_to_excel(input_path, output_path)
            
            elif input_extension in ["pptx", "ppt"] and output_format == "pdf":
                return self._convert_ppt_to_pdf(input_path, output_path)
            
            else:
                self.logger.error(f"Conversion from {input_extension} to {output_format} is not supported")
                return False
                
        except Exception as e:
            self.logger.error(f"Error converting {input_path}: {str(e)}")
            return False
    
    def _convert_pdf_to_docx(self, input_path, output_path):
        """Convert PDF to DOCX using pdf2docx"""
        try:
            cv = Converter(input_path)
            cv.convert(output_path)
            cv.close()
            return True
        except Exception as e:
            self.logger.error(f"Error converting PDF to DOCX: {str(e)}")
            return False
    
    def _convert_pdf_to_txt(self, input_path, output_path):
        """Convert PDF to TXT using PyPDF2"""
        try:
            with open(input_path, 'rb') as file:
                reader = PyPDF2.PdfFileReader(file)
                with open(output_path, 'w', encoding='utf-8') as output_file:
                    for page_num in range(reader.numPages):
                        text = reader.getPage(page_num).extractText()
                        output_file.write(text)
            return True
        except Exception as e:
            self.logger.error(f"Error converting PDF to TXT: {str(e)}")
            return False
    
    def _convert_doc_to_pdf(self, input_path, output_path):
        """Convert DOC/DOCX to PDF"""
        try:
            # This typically requires external tools like LibreOffice
            # For a complete implementation, you would need to install LibreOffice
            # and use subprocess to call it for conversion
            if os.name == 'posix':  # Linux/Mac
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(output_path), input_path
                ]
            else:  # Windows
                cmd = [
                    'soffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(output_path), input_path
                ]
            
            process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            stdout, stderr = process.communicate()
            
            if process.returncode != 0:
                self.logger.error(f"Error converting DOC to PDF: {stderr.decode()}")
                return False
                
            # The output file from LibreOffice will have the same name as input but with .pdf extension
            # If that's not what we want, we need to rename it
            libreoffice_output = os.path.join(
                os.path.dirname(output_path),
                os.path.splitext(os.path.basename(input_path))[0] + '.pdf'
            )
            
            if libreoffice_output != output_path and os.path.exists(libreoffice_output):
                os.rename(libreoffice_output, output_path)
                
            return True
            
        except Exception as e:
            self.logger.error(f"Error converting DOC to PDF: {str(e)}")
            return False
    
    def _convert_doc_to_txt(self, input_path, output_path):
        """Convert DOC/DOCX to TXT"""
        try:
            doc = Document(input_path)
            with open(output_path, 'w', encoding='utf-8') as f:
                for paragraph in doc.paragraphs:
                    f.write(paragraph.text + '\n')
            return True
        except Exception as e:
            self.logger.error(f"Error converting DOC to TXT: {str(e)}")
            return False
    
    def _convert_txt_to_pdf(self, input_path, output_path):
        """Convert TXT to PDF using reportlab"""
        try:
            # For a complete implementation, you would need to install reportlab
            # and use it to create a PDF from the text file
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            with open(input_path, 'r', encoding='utf-8') as f:
                text = f.read()
                para = Paragraph(text.replace('\n', '<br/>'), styles['Normal'])
                story.append(para)
                
            doc.build(story)
            return True
            
        except Exception as e:
            self.logger.error(f"Error converting TXT to PDF: {str(e)}")
            return False
    
    def _convert_txt_to_docx(self, input_path, output_path):
        """Convert TXT to DOCX"""
        try:
            doc = Document()
            
            with open(input_path, 'r', encoding='utf-8') as f:
                for line in f:
                    doc.add_paragraph(line.strip())
                    
            doc.save(output_path)
            return True
            
        except Exception as e:
            self.logger.error(f"Error converting TXT to DOCX: {str(e)}")
            return False
    
    def _convert_excel_to_csv(self, input_path, output_path):
        """Convert Excel to CSV"""
        try:
            df = pd.read_excel(input_path)
            df.to_csv(output_path, index=False)
            return True
        except Exception as e:
            self.logger.error(f"Error converting Excel to CSV: {str(e)}")
            return False
    
    def _convert_csv_to_excel(self, input_path, output_path):
        """Convert CSV to Excel"""
        try:
            df = pd.read_csv(input_path)
            df.to_excel(output_path, index=False)
            return True
        except Exception as e:
            self.logger.error(f"Error converting CSV to Excel: {str(e)}")
            return False
    
    def _convert_ppt_to_pdf(self, input_path, output_path):
        """Convert PowerPoint to PDF"""
        try:
            # This typically requires external tools like LibreOffice
            # Similar to _convert_doc_to_pdf
            if os.name == 'posix':  # Linux/Mac
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(output_path), input_path
                ]
            else:  # Windows
                cmd = [
                    'soffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(output_path), input_path
                ]
            
            process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            stdout, stderr = process.communicate()
            
            if process.returncode != 0:
                self.logger.error(f"Error converting PPT to PDF: {stderr.decode()}")
                return False
                
            # The output file from LibreOffice will have the same name as input but with .pdf extension
            # If that's not what we want, we need to rename it
            libreoffice_output = os.path.join(
                os.path.dirname(output_path),
                os.path.splitext(os.path.basename(input_path))[0] + '.pdf'
            )
            
            if libreoffice_output != output_path and os.path.exists(libreoffice_output):
                os.rename(libreoffice_output, output_path)
                
            return True
            
        except Exception as e:
            self.logger.error(f"Error converting PPT to PDF: {str(e)}")
            return False
    
    @staticmethod
    def is_supported_input(file_path):
        """Check if a file is a supported input format"""
        extension = os.path.splitext(file_path)[1][1:].lower()
        return extension in DocumentConverter.SUPPORTED_INPUT_FORMATS